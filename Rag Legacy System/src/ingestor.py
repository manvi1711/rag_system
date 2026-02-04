# Multimodal ingestion (PDF, TXT, DOCX, images + OCR)
import os
import numpy as np
import easyocr
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .bedrock_embedder import TitanEmbedder
from src.config import REGION_NAME, BEDROCK_EMBED_MODEL, CHUNK_SIZE, CHUNK_OVERLAP
from langchain_community.vectorstores.faiss import FAISS
from langchain_core.documents import Document
from langchain_community.docstore.in_memory import InMemoryDocstore


def chunk_text_with_metadata(text, source_name):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )

    raw_chunks = splitter.split_text(text)

    chunks = []
    for i, chunk in enumerate(raw_chunks):
        chunks.append({
            "content": chunk,
            "source": source_name,
            "page": i
        })

    return chunks


def load_documents():
    from src.config import DOCUMENT_PATH
    reader_ocr = easyocr.Reader(['en'])

    print(f"Loading documents from: {DOCUMENT_PATH}")

    if not os.path.exists(DOCUMENT_PATH):
        raise FileNotFoundError(f"Document path does not exist: {DOCUMENT_PATH}")

    all_files = [
        f for f in os.listdir(DOCUMENT_PATH)
        if f.lower().endswith((".pdf", ".txt", ".docx", ".png", ".jpg", ".jpeg"))
    ]

    if not all_files:
        raise ValueError(f"No supported files found in {DOCUMENT_PATH}")

    all_chunks = []

    for filename in all_files:
        file_path = os.path.join(DOCUMENT_PATH, filename)
        print(f"\n  Loading: {filename}")

        if filename.lower().endswith(".pdf"):
            reader = PdfReader(file_path)

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if not page_text:
                    continue

                page_chunks = chunk_text_with_metadata(page_text, filename)
                for c in page_chunks:
                    c["page"] = page_num

                all_chunks.extend(page_chunks)

            print(f"   Loaded {len(reader.pages)} pages.")

        elif filename.lower().endswith(".txt"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

            txt_chunks = chunk_text_with_metadata(text, filename)
            for c in txt_chunks:
                c["page"] = 0

            all_chunks.extend(txt_chunks)
            print(f"   Loaded TXT with {len(txt_chunks)} chunks.")

        elif filename.lower().endswith(".docx"):
            doc = DocxDocument(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            docx_chunks = chunk_text_with_metadata(full_text, filename)
            for c in docx_chunks:
                c["page"] = 0

            all_chunks.extend(docx_chunks)
            print(f"   Loaded DOCX with {len(docx_chunks)} chunks.")

        elif filename.lower().endswith((".png", ".jpg", ".jpeg")):
            ocr_text = reader_ocr.readtext(file_path, detail=0)
            extracted_text = "\n".join(ocr_text)

            img_chunks = chunk_text_with_metadata(extracted_text, filename)
            for c in img_chunks:
                c["page"] = 0

            all_chunks.extend(img_chunks)
            print(f"   Loaded image (OCR) with {len(img_chunks)} chunks.")

    if not all_chunks:
        raise ValueError("No extractable text found.")

    print(f"\nSuccessfully created {len(all_chunks)} chunks.")
    return all_chunks


def build_faiss_vectorstore(chunks):
    from langchain_core.documents import Document
    
    embedder = TitanEmbedder(BEDROCK_EMBED_MODEL, REGION_NAME)
    
    documents = [
        Document(
            page_content=c["content"],
            metadata={"source": c["source"], "page": c["page"]}
        )
        for c in chunks
    ]
    
    print("Building vectorstore with embeddings...")
    
    vectorstore = FAISS.from_documents(
        documents=documents,
        embedding=embedder
    )
    
    return vectorstore

def ingest_document(_unused):
    print("\n===== STARTING INGESTION PIPELINE =====")

    chunks = load_documents()

    vectorstore = build_faiss_vectorstore(chunks)

    print("Saving FAISS index...")
    vectorstore.save_local("faiss_index")

    print("===== INGESTION COMPLETED SUCCESSFULLY =====")
